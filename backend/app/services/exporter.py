from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def make_docx(title: str, content: str) -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(1.25)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    for block in content.split("\n"):
        paragraph = doc.add_paragraph(block)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
